#!/usr/bin/env python3
"""
MD2DOCX_ND30 - Chuyển đổi Markdown sang Word (.docx) chuẩn Nghị định 30/2020/NĐ-CP
==================================================================================
Quy chuẩn văn bản hành chính nhà nước (Nghị định 30/2020/NĐ-CP):
- Khổ giấy: A4 (210 x 297 mm)
- Phông chữ: Times New Roman (Unicode TCVN 6909:2001)
- Lề trang: Trên 2.0 cm, Dưới 2.0 cm, Trái 3.0 cm, Phải 2.0 cm
- Nội dung: Cỡ chữ 13-14 pt, căn đều 2 bên (Justified), lùi đầu dòng 1.0 - 1.27 cm, giãn dòng 1.3 - 1.5.
- Tiêu đề mục (Headings): Cỡ chữ 13-15 pt, in hoa/in đậm đúng quy chuẩn.
- Bảng biểu: Căn giữa, viền 0.5pt, tiêu đề dòng đậm/nền xám nhẹ.
"""

import os
import sys
import argparse
import re
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# --- Cấu hình mặc định theo Nghị định 30/2020/NĐ-CP ---
DEFAULT_FONT_NAME = "Times New Roman"
DEFAULT_BODY_SIZE = 13  # 13 pt
DEFAULT_LINE_SPACING = 1.3
DEFAULT_MARGIN_TOP = 2.0    # 2.0 cm
DEFAULT_MARGIN_BOTTOM = 2.0 # 2.0 cm
DEFAULT_MARGIN_LEFT = 3.0   # 3.0 cm (Lề trái rộng để đóng sổ)
DEFAULT_MARGIN_RIGHT = 2.0  # 2.0 cm
DEFAULT_FIRST_LINE_INDENT = 1.0 # 1.0 cm

def setup_document_geometry(doc, top=2.0, bottom=2.0, left=3.0, right=2.0):
    """Cấu hình khổ giấy A4 và lề trang chuẩn Nghị định 30/2020/NĐ-CP."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)

def set_cell_background(cell, hex_color):
    """Đặt màu nền cho ô trong bảng."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color="B0C4DE", sz="4", val="single"):
    """Đặt đường viền nét đơn mảnh cho toàn bộ bảng."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Cấu hình đệm nội dung ô (padding) giúp bảng thông thoáng."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>\n'
        f'  <w:top w:w="{top}" w:type="dxa"/>\n'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
        f'  <w:left w:w="{left}" w:type="dxa"/>\n'
        f'  <w:right w:w="{right}" w:type="dxa"/>\n'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_header_footer_page_number(doc):
    """Thêm đánh số trang ở giữa phần Header/Footer theo phông Times New Roman 13pt."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.font.name = DEFAULT_FONT_NAME
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Thêm trường PAGE động của Word
        fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
        fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
        fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        p._p.append(fldChar1)
        p._p.append(instrText)
        p._p.append(fldChar2)
        p._p.append(fldChar3)

class MarkdownToDocxND30:
    def __init__(self, font_name=DEFAULT_FONT_NAME, body_size=DEFAULT_BODY_SIZE, line_spacing=DEFAULT_LINE_SPACING):
        self.font_name = font_name
        self.body_size = body_size
        self.line_spacing = line_spacing

    def parse_element_to_runs(self, element, paragraph, base_bold=False, base_italic=False, base_color=None, is_code=False):
        """Duyệt qua các thẻ inline HTML (b, i, code, a, text) và tạo Run cho Word."""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if not text:
                    continue
                run = paragraph.add_run(text)
                run.font.name = self.font_name
                run.font.size = Pt(self.body_size)
                run.bold = base_bold
                run.italic = base_italic
                if is_code:
                    run.font.name = "Consolas"
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(180, 40, 40)
                elif base_color:
                    run.font.color.rgb = base_color
            elif isinstance(child, Tag):
                tag_name = child.name.lower()
                bold = base_bold or (tag_name in ['strong', 'b'])
                italic = base_italic or (tag_name in ['em', 'i'])
                code = is_code or (tag_name == 'code')
                color = base_color
                
                if tag_name == 'a':
                    color = RGBColor(0, 80, 160)
                    
                self.parse_element_to_runs(child, paragraph, base_bold=bold, base_italic=italic, base_color=color, is_code=code)

    def process_heading(self, doc, tag):
        """Xử lý các tiêu đề H1-H6 theo chuẩn văn bản hành chính."""
        level = int(tag.name[1])
        text = tag.get_text().strip()
        
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.first_line_indent = Cm(0) # Tiêu đề không lùi đầu dòng
        
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(text.upper()) # In hoa H1
            run.font.name = self.font_name
            run.font.size = Pt(15)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue sang trọng
        elif level == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.font.name = self.font_name
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 3:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.font.name = self.font_name
            run.font.size = Pt(13)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(30, 30, 30)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.name = self.font_name
            run.font.size = Pt(13)
            run.italic = True

    def process_paragraph(self, doc, tag):
        """Xử lý đoạn văn thường (Paragraph) - căn đều 2 bên, lùi đầu dòng 1.0cm."""
        text = tag.get_text().strip()
        if not text:
            return
            
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = self.line_spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Cm(DEFAULT_FIRST_LINE_INDENT)
        
        self.parse_element_to_runs(tag, p)

    def process_list(self, doc, tag, level=0):
        """Xử lý danh sách có thứ tự hoặc không thứ tự."""
        is_ordered = (tag.name == 'ol')
        
        for idx, li in enumerate(tag.find_all('li', recursive=False), 1):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = self.line_spacing
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(1.0 + level * 0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            
            # Ký hiệu bullet hoặc số
            prefix = f"{idx}. " if is_ordered else "• "
            run_prefix = p.add_run(prefix)
            run_prefix.font.name = self.font_name
            run_prefix.font.size = Pt(self.body_size)
            run_prefix.bold = True
            
            self.parse_element_to_runs(li, p)
            
            # Duyệt các danh sách con lồng nhau
            for sub_list in li.find_all(['ul', 'ol'], recursive=False):
                self.process_list(doc, sub_list, level=level+1)

    def process_blockquote(self, doc, tag):
        """Xử lý khối trích dẫn / Ghi chú (Blockquote / Alert)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(0)
        
        self.parse_element_to_runs(tag, p, base_italic=True, base_color=RGBColor(80, 80, 80))

    def process_codeblock(self, doc, tag):
        """Xử lý khối mã nguồn (Fenced Code Block)."""
        code_text = tag.get_text()
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.1
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(0)
        
        run = p.add_run(code_text)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(40, 40, 40)

    def process_table(self, doc, tag):
        """Xử lý bảng dữ liệu (Table) chuẩn khung viền Nghị định 30."""
        rows = tag.find_all('tr')
        if not rows:
            return
            
        num_rows = len(rows)
        num_cols = max(len(r.find_all(['td', 'th'])) for r in rows)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color="808080", sz="4", val="single")
        
        for r_idx, tr in enumerate(rows):
            cells = tr.find_all(['td', 'th'])
            for c_idx, cell_tag in enumerate(cells):
                if c_idx >= num_cols:
                    break
                cell = table.cell(r_idx, c_idx)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.first_line_indent = Cm(0)
                
                is_header = (r_idx == 0) or (cell_tag.name == 'th')
                if is_header:
                    set_cell_background(cell, "F2F4F7") # Nền xám nhạt tiêu đề
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.parse_element_to_runs(cell_tag, p, base_bold=True)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    self.parse_element_to_runs(cell_tag, p)

    def convert(self, md_text: str, output_docx_path: str):
        """Chuyển đổi chuỗi Markdown thành file Word (.docx)."""
        # Render Markdown sang HTML bằng các extensions hỗ trợ bảng, codeblock
        html = markdown.markdown(
            md_text,
            extensions=['tables', 'fenced_code', 'nl2br', 'sane_lists']
        )
        soup = BeautifulSoup(html, 'html.parser')
        
        doc = Document()
        setup_document_geometry(doc, top=self.margin_top, bottom=self.margin_bottom, left=self.margin_left, right=self.margin_right)
        add_header_footer_page_number(doc)
        
        # Duyệt các thẻ HTML cấp cao nhất
        for tag in soup.children:
            if isinstance(tag, NavigableString):
                continue
            tag_name = tag.name.lower()
            
            if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self.process_heading(doc, tag)
            elif tag_name == 'p':
                self.process_paragraph(doc, tag)
            elif tag_name in ['ul', 'ol']:
                self.process_list(doc, tag)
            elif tag_name == 'blockquote':
                self.process_blockquote(doc, tag)
            elif tag_name == 'pre':
                self.process_codeblock(doc, tag)
            elif tag_name == 'table':
                self.process_table(doc, tag)
            elif tag_name == 'hr':
                doc.add_page_break()

        doc.save(output_docx_path)
        print(f"[✅ THÀNH CÔNG] Đã chuyển đổi: {output_docx_path}")

def convert_md_to_docx(input_path: str, output_path: str = None, font_size=13, line_spacing=1.3, margin_top=2.0, margin_bottom=2.0, margin_left=3.0, margin_right=2.0):
    input_p = Path(input_path)
    if not input_p.exists():
        print(f"[-] Lỗi: Không tìm thấy tệp đầu vào {input_path}")
        sys.exit(1)

    converter = MarkdownToDocxND30(body_size=font_size, line_spacing=line_spacing)
    converter.margin_top = margin_top
    converter.margin_bottom = margin_bottom
    converter.margin_left = margin_left
    converter.margin_right = margin_right

    if input_p.is_file():
        if not output_path:
            output_path = str(input_p.with_suffix('.docx'))
        with open(input_p, 'r', encoding='utf-8') as f:
            md_text = f.read()
        converter.convert(md_text, output_path)
    elif input_p.is_dir():
        out_dir = Path(output_path) if output_path else input_p
        out_dir.mkdir(parents=True, exist_ok=True)
        md_files = list(input_p.glob('**/*.md'))
        print(f"[+] Tìm thấy {len(md_files)} tệp Markdown trong thư mục {input_path}")
        for md_file in md_files:
            rel_p = md_file.relative_to(input_p)
            out_file = out_dir / rel_p.with_suffix('.docx')
            out_file.parent.mkdir(parents=True, exist_ok=True)
            with open(md_file, 'r', encoding='utf-8') as f:
                md_text = f.read()
            converter.convert(md_text, str(out_file))

def main():
    parser = argparse.ArgumentParser(description="Convert Markdown files to Word (.docx) complying with VN State Decree 30/2020/NĐ-CP.")
    parser.add_argument("input", help="Path to input .md file or directory containing .md files")
    parser.add_argument("-o", "--output", help="Path to output .docx file or output directory")
    parser.add_argument("--font-size", type=int, default=13, help="Body text font size (default: 13pt)")
    parser.add_argument("--line-spacing", type=float, default=1.3, help="Line spacing factor (default: 1.3)")
    parser.add_argument("--margin-top", type=float, default=2.0, help="Top margin in cm (default: 2.0)")
    parser.add_argument("--margin-bottom", type=float, default=2.0, help="Bottom margin in cm (default: 2.0)")
    parser.add_argument("--margin-left", type=float, default=3.0, help="Left margin in cm (default: 3.0)")
    parser.add_argument("--margin-right", type=float, default=2.0, help="Right margin in cm (default: 2.0)")
    
    args = parser.parse_args()
    convert_md_to_docx(
        args.input,
        args.output,
        font_size=args.font_size,
        line_spacing=args.line_spacing,
        margin_top=args.margin_top,
        margin_bottom=args.margin_bottom,
        margin_left=args.margin_left,
        margin_right=args.margin_right
    )

if __name__ == "__main__":
    main()
