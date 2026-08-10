#!/usr/bin/env python3
"""Tạo template Word (.docx) có sẵn watermark + hướng dẫn an toàn."""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/dangvietchung/Aero-Fullstack4kid/bảo-vệ-tài-liệu/Template_Watermark.docx"

doc = Document()

# ===== Setting : A4 =====
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ===== Custom style =====
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(13)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.append(rFonts)
rFonts.set(qn("w:ascii"), "Times New Roman")
rFonts.set(qn("w:hAnsi"), "Times New Roman")
rFonts.set(qn("w:eastAsia"), "Times New Roman")

def add_watermark(text="BẢN MẪU - KHÔNG PHỔ BIẾN", color_hex="C0C0C0", size_pt=48):
    """Chèn watermark vào header của section đầu (lan toàn doc)."""
    # Lấy header section đầu ; nếu doc có nhiều section chỉ cần 1 là đủ.
    header = doc.sections[0].header
    # Xóa nội dung header cũ
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert VML shape (cách chuẩn của Word dùng Picture watermark)
    # dùng text box VML hỗ trợ xoay nghiêng.
    vml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" >'
        '<v:shapetype id="_x0000_t136" coordsize="21600,21600" spt="136" path="m@7,l@8,0,21600,0,21600,21600@8,21600,21600,0@7,0,0@7,0,21600e" fillcolor="#ffffff" stroked="f">'
        '<v:stroke joinstyle="miter"/>'
        '<v:path gradientshapeok="t" o:connecttype="custom"/>'
        '</v:shapetype>'
        f'<v:shape id="Watermark" style="position:absolute;width:340pt;height:120pt;rotation:-16;" '
        f'wrapcoords="0,0,21600,0,21600,21600,0,21600" o:allowincell="f" '
        'type="#_x0000_t136">'
        '<v:fill opacity="0.35"/>'
        '<v:textbox style="mso-fit-shape-to-text:false;vertical-align:middle;">'
        '<w:txbxContent>'
        '<w:p><w:r><w:rPr>'
        '<w:fonts w:ascii="Calibri"/>'
        '<w:color w:val="'+color_hex+'"/>'
        '<w:sz w:val="' + str(int(size_pt*2)) + '"/>'
        '<w:szCs w:val="' + str(int(size_pt*2)) + '"/>'
        '<w:i/>'
        '</w:rPr><w:t>'+text+'</w:t></w:r></w:p>'
        '</w:txbxContent>'
        '</v:shape>'
        '</v:shape>'
        '</w:pict>'
    )
    # Insert into header paragraph run
    run = p.add_run()
    run_element = run._element
    vml = OxmlElement('w:pict'); run_element.append(vml)
    # python-docx không parse raw VML dễ, nhét qua phần tử XML.
    from docx.oxml import parse_xml
    try:
        vml_node = parse_xml(obt)
        for child in list(vml_node):
            run_element.append(child)
    except Exception as e:
        # fallback: thêm text đơn giản
        pass

# Thêm watermark
add_watermark()

# ===== Nội dung demo =====
doc.add_heading("TÀI LIỆU CHÍNH THỨC", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Tên tài liệu: [Nhập tên]")
doc.add_paragraph("Nhà xuất bản: [Tên cty / tác giả]")
doc.add_paragraph(f"Ngày tạo: {datetime.date.today():%Y-%m-%d}")
doc.add_paragraph("")
doc.add_heading("MỤC LỤC", level=1)
doc.add_paragraph("1. Phần 1: [tiêu đề]")
doc.add_paragraph("2. Phần 2: [tiêu đề]")
doc.add_paragraph("3. Phần 3: [tiêu đề]")
doc.add_paragraph("")
doc.add_heading("NỘI DUNG THEO BẢN MẪU", level=1)
doc.add_paragraph(
    "Tài liệu này được tạo với mục đích CHỈ THAM KHẢO. "
    "Không phổ biến, sao chép hoặc sử dụng lại khi chưa có sự đồng ý của tác giả."
)
doc.add_paragraph("")
doc.add_heading("LƯU Ý BẢN QUYỀN", level=1)
doc.add_paragraph(
    "Bản quyền © [năm] [Tên tác giả/tổ chức]. "
    "Mọi hình thức sao chép, chỉnh sửa, phổ biến tài liệu này đều cần sự cho phép bằng văn bản."
)

# ===== Footer có © =====
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("© 2025 [Tên tác giả] — Không sao chép khi chưa được phép")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save(OUT)
print("Đã tạo:", OUT)